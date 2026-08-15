"""
DOCX Source Adapter for ingesting Microsoft Word technical documents using python-docx.
"""

import os
import uuid
from pathlib import Path
from docx import Document

from backend.ingestion.base import BaseSourceAdapter, IngestionError
from backend.ingestion.models import SourceDocument


class DocxSourceAdapter(BaseSourceAdapter):
    """Adapter for ingesting DOCX product documentation."""

    def ingest(self, file_path: str) -> SourceDocument:
        if not file_path or not os.path.exists(file_path):
            raise IngestionError("DOCX file does not exist on disk.")

        filename = Path(file_path).name
        source_id = f"docx_{uuid.uuid4().hex[:12]}"

        try:
            doc = Document(file_path)
            parts = []

            for p in doc.paragraphs:
                p_text = p.text.strip()
                if p_text:
                    if p.style.name.startswith("Heading"):
                        parts.append(f"\nHEADING ({p.style.name}): {p_text}")
                    else:
                        parts.append(p_text)

            for t_idx, table in enumerate(doc.tables, start=1):
                table_rows = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        table_rows.append(" | ".join(row_cells))
                if table_rows:
                    parts.append(f"\nTABLE {t_idx}:\n" + "\n".join(table_rows))

        except Exception as e:
            raise IngestionError(f"Failed to parse DOCX file: {str(e)}") from e

        if not parts:
            raise IngestionError("DOCX document contains no extractable text.")

        full_content = f"DOCX DOCUMENT: {filename}\n\n" + "\n".join(parts)

        return SourceDocument(
            source_id=source_id,
            source_type="docx",
            source_name=filename,
            content=full_content,
            local_path=file_path,
            metadata={
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            },
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
